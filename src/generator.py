import os
import re
from docx import Document

def replace_text_in_paragraphs(doc, replacements):
    """
    Scans all paragraphs AND table cells in the document and replaces standard text 
    placeholders like {{ client_name }} with the actual values.
    """
    pattern = re.compile(r'\{\{\s*(.*?)\s*\}\}')
    missing_keys = set()

    # 1. Scan and Replace in normal text paragraphs
    for paragraph in doc.paragraphs:
        matches = pattern.findall(paragraph.text)
        for key in matches:
            if key not in replacements or replacements[key] == "":
                missing_keys.add(key)
        
        for key, value in replacements.items():
            placeholder = f"{{{{ {key} }}}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                
    # 2. Scan and Replace inside table cells (Crucial for physical spec tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = pattern.findall(paragraph.text)
                    for key in matches:
                        if key not in replacements or replacements[key] == "":
                            missing_keys.add(key)

                    for key, value in replacements.items():
                        placeholder = f"{{{{ {key} }}}}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))

    # 3. Terminal Report
    if missing_keys:
        print("\n⚠️  DEBUG WARNING: Placeholders missing data payload:")
        for key in missing_keys:
            print(f"   -> Couldn't find data for: '{key}'")
        print("   (These will be left as raw brackets in the final document)\n")


def generate_clearit_doc(extracted_data, template_path, output_filename):
    print("\n" + "="*50)
    print("🚀 RUNNING FINALIZED WORD GENERATION ENGINE")
    print("="*50)

    # --- STEP 1: Flatten Automated Data Arrays into Layout Placeholders ---
    
    # Flatten Network Information array into individual numbered cells
    net_info = extracted_data.get("network_information", [])
    if net_info:
        # Sort or safely map by host index matching the template order
        for idx, node in enumerate(net_info, start=1):
            extracted_data[f"cvm_ip_{idx}"] = node.get("cvm_ip", "")
            extracted_data[f"ahv_ip_{idx}"] = node.get("mgmt_ip", "")
            extracted_data[f"ipmi_ip_{idx}"] = node.get("ipmi_ip", "")
            
    # Pull global cluster metadata safely into flat table strings
    prism_el = extracted_data.get("prism_element", [])
    if prism_el:
        extracted_data["prism_element_ip"] = prism_el[0].get("cluster_ip", "")
        extracted_data["cluster_name"] = prism_el[0].get("cluster_name", "")
        extracted_data["cluster_rf"] = prism_el[0].get("cluster_rf", "")

    # Pull global service names safely
    srv_info = extracted_data.get("service_name_hour", [])
    if srv_info:
        extracted_data["dns_servers"] = srv_info[0].get("dns_server", "")

    # Map host details into sequential layout lists safely 
    host_details = extracted_data.get("hosts_details", [])
    host_models = extracted_data.get("hosts_models", [])
    for idx in range(len(host_details)):
        if idx < len(host_details):
            extracted_data[f"node_{idx+1}_hostname"] = host_details[idx].get("hostname", "")
        if idx < len(host_models):
            extracted_data[f"node_serial_{idx+1}"] = host_models[idx].get("serial", "")

    # Load and process the document exactly like before
    doc = Document(template_path)
    
    print("📝 Injecting text placeholders and flat technical parameters...")
    replace_text_in_paragraphs(doc, extracted_data)

    # --- B: Populate Dynamic Lists ---
    print("📊 3. Dynamically building tables...")
    for idx, table in enumerate(doc.tables):
        if not table.rows:
            continue
        
        header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
        
        # 1. Hosts Details Table
        if "Hostname" in header and "Hypervisor Version" in header:
            for item in extracted_data.get("hosts_details", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("hostname", ""))
                cells[1].text = str(item.get("hypervisor_version", ""))
                
        # 2. Hosts Models Table
        elif "Hostname" in header and "Model" in header and "Serial" in header:
            for item in extracted_data.get("hosts_models", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("hostname", ""))
                cells[1].text = str(item.get("model", ""))
                cells[2].text = str(item.get("serial", ""))
                
        # 3. Hosts Information (CPU/Memory)
        elif "Hostname" in header and "CPU Cores" in header and "Memory" in header:
            for item in extracted_data.get("hosts_informations", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("hostname", ""))
                cells[1].text = str(item.get("cpu_cores", ""))
                cells[2].text = str(item.get("cpu_threads", ""))
                cells[3].text = str(item.get("memory", ""))
                
        # 4. BIOS & BMC
        elif "Hostname" in header and "Bios Version" in header and "BMC Version" in header:
            for item in extracted_data.get("bios_bmc", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("hostname", ""))
                cells[1].text = str(item.get("bios_version", ""))
                cells[2].text = str(item.get("bios_model", ""))
                cells[3].text = str(item.get("bmc_version", ""))
                cells[4].text = str(item.get("bmc_model", ""))
                
        # 5. Network Information
        elif "Hostname" in header and "CVM IP" in header and "Management IP" in header:
            for item in extracted_data.get("network_information", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("hostname", ""))
                cells[1].text = str(item.get("cvm_ip", ""))
                cells[2].text = str(item.get("mgmt_ip", ""))
                cells[3].text = str(item.get("ipmi_ip", ""))
                
        # 6. NTP / DNS
        elif "NTP Server(s)" in header and "DNS Server(s)" in header:
            for item in extracted_data.get("service_name_hour", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("ntp_server", ""))
                cells[1].text = str(item.get("dns_server", ""))
                cells[2].text = str(item.get("global_whitelist", ""))
                
        # 7. Storage Pools
        elif "Name" in header and "Storage Pool ID" in header and "Max Capacity" in header:
            for item in extracted_data.get("storage_pools", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("storage_pool_name", ""))
                cells[1].text = str(item.get("storage_pool_id", ""))
                cells[2].text = str(item.get("max_capacity", ""))
                cells[3].text = str(item.get("ilm_threshold", ""))

        # 8. Containers List
        elif "Container Name" in header and "Max Usable Capacity" in header:
            for item in extracted_data.get("containers_list", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("hostname", ""))
                cells[1].text = str(item.get("max_usable_capacity", ""))
                cells[2].text = str(item.get("total_raw_capacity", ""))
                cells[3].text = str(item.get("reserved_capacity", ""))
                cells[4].text = str(item.get("nfs_whitelist", ""))

        # 9. Containers Options
        elif "Container Name" in header and "Compression Enabled" in header:
            for item in extracted_data.get("containers_options", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("hostname", ""))
                cells[1].text = str(item.get("rf", ""))
                cells[2].text = str(item.get("compression_enabled", ""))
                cells[3].text = str(item.get("compression_delay", ""))
                cells[4].text = str(item.get("ssd_dedup", ""))
                cells[5].text = str(item.get("hdd_dedup", ""))
                cells[6].text = str(item.get("erasure_coding", ""))

        # 10. Control Virtual Machines (CVM)
        elif "CVM" in header and "Memory" in header and "vCPU" in header:
            for item in extracted_data.get("control_virtual_machine", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("cvm_name", ""))
                cells[1].text = str(item.get("memory", ""))
                cells[2].text = str(item.get("vcpu", ""))
                cells[3].text = str(item.get("ip_address", ""))

        # 11. Licensing
        elif "License" in header and "License Type" in header and "Block Serial Number" in header:
            for item in extracted_data.get("licensing", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("license_name", ""))
                cells[1].text = str(item.get("license_type", ""))
                cells[2].text = str(item.get("block_serial_number", ""))
                cells[3].text = str(item.get("expiration", ""))

        # 12. Alert Monitoring
        elif "Host Name" in header and "Port" in header and "Security Mode" in header:
            for item in extracted_data.get("alert_monitoring", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("host_name", ""))
                cells[1].text = str(item.get("port", ""))
                cells[2].text = str(item.get("security_mode", ""))
                cells[3].text = str(item.get("username", ""))
                cells[4].text = str(item.get("email_address", ""))

        # 13. Directory List
        elif "Directory Type" in header and "Connection Type" in header and "Domain" in header:
            for item in extracted_data.get("list_directory", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("directory_name", ""))
                cells[1].text = str(item.get("directory_type", ""))
                cells[2].text = str(item.get("connection_type", ""))
                cells[3].text = str(item.get("directory_url", ""))
                cells[4].text = str(item.get("directory_domain", ""))

        # 14. Prism Element
        elif "Cluster Name" in header and "Cluster UUID" in header and "Cluster IP" in header:
            for item in extracted_data.get("prism_element", []):
                cells = table.add_row().cells
                cells[0].text = str(item.get("cluster_name", ""))
                cells[1].text = str(item.get("cluster_uuid", ""))
                cells[2].text = str(item.get("cluster_ip", ""))
                cells[3].text = str(item.get("cluster_rf", ""))

    # --- C: Save Output ---
    output_dir = "data/generated_docs"
    os.makedirs(output_dir, exist_ok=True)
    final_path = os.path.join(output_dir, output_filename)
    
    print(f"💾 4. Saving generated document to: {final_path}")
    doc.save(final_path)
    
    print("="*50)
    print("🎉 GENERATION ENGINE FINISHED")
    print("="*50 + "\n")
    
    return final_path